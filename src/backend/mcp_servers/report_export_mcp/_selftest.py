"""Selftest for report_export_mcp.

Run:
  python -m mcp_servers.report_export_mcp._selftest
"""

from __future__ import annotations

import importlib
import inspect
import io
import sys
from unittest.mock import patch


def _ok(msg: str) -> None:
    print(f"[OK] {msg}")


def _fail(msg: str) -> None:
    print(f"[FAIL] {msg}", file=sys.stderr)
    raise SystemExit(1)


_SAMPLE_MD = """\
## 测试标题

这是**加粗文字**和*斜体文字*以及`行内代码`。

[链接文字](https://example.com)

| 姓名 | 年龄 | 城市 |
|------|------|------|
| 张三 | 25   | 北京 |
| 李四 | 30   | 上海 |

```python
def hello():
    print("world")
```

1. 第一项
2. 第二项

- 无序一
- 无序二
"""


def _test_import():
    try:
        mod = importlib.import_module("mcp_servers.report_export_mcp.server")
        _ok("import server")
    except ModuleNotFoundError as e:
        _ok(f"SKIP: import server (missing dependency): {e}")
        print("SELFTEST_SKIP")
        return None
    except Exception as e:
        _fail(f"import server failed: {e!r}")

    fn = getattr(mod, "export_report_to_docx", None)
    if fn is None:
        _fail("missing tool function: export_report_to_docx")
    sig = inspect.signature(fn)
    _ok(f"signature: export_report_to_docx{sig}")

    impl = importlib.import_module("mcp_servers.report_export_mcp.impl")
    if getattr(impl, "export_report_to_docx", None) is None:
        _fail("missing impl: export_report_to_docx")
    _ok("import impl")
    return impl


def _load_docx(docx_bytes: bytes):
    from docx import Document
    return Document(io.BytesIO(docx_bytes))


def _test_bold_italic(impl):
    """Verify **bold** and *italic* produce correct run formatting."""
    raw = impl._fallback_markdown_to_docx("**加粗** 和 *斜体*", "test")
    doc = _load_docx(raw)

    found_bold = False
    found_italic = False
    for para in doc.paragraphs:
        for run in para.runs:
            if run.bold and "加粗" in run.text:
                found_bold = True
            if run.italic and "斜体" in run.text:
                found_italic = True

    if not found_bold:
        _fail("bold not rendered in fallback")
    if not found_italic:
        _fail("italic not rendered in fallback")
    _ok("bold + italic (fallback)")


def _test_table_render(impl):
    """Verify markdown table → doc.tables with correct dimensions."""
    md = "| A | B |\n|---|---|\n| 1 | 2 |\n| 3 | 4 |"
    raw = impl._fallback_markdown_to_docx(md, "table-test")
    doc = _load_docx(raw)

    if not doc.tables:
        _fail("no tables found in fallback output")
    t = doc.tables[0]
    if len(t.rows) != 3:
        _fail(f"expected 3 rows (header+2 data), got {len(t.rows)}")
    if len(t.columns) != 2:
        _fail(f"expected 2 columns, got {len(t.columns)}")
    _ok("table render (fallback): 3 rows x 2 cols")


def _test_heading_font(impl):
    """Verify heading paragraphs get the heading CJK font applied."""
    raw = impl._fallback_markdown_to_docx("## 标题二", "heading-test")
    doc = _load_docx(raw)

    from docx.oxml.ns import qn
    found = False
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                rPr = pPr.find(qn("w:rPr"))
                if rPr is not None:
                    rFonts = rPr.find(qn("w:rFonts"))
                    if rFonts is not None:
                        ea = rFonts.get(qn("w:eastAsia"))
                        if ea == impl._HEADING_FONT:
                            found = True

    if not found:
        _fail(f"heading font not set to {impl._HEADING_FONT}")
    _ok(f"heading font = {impl._HEADING_FONT}")


def _test_fallback_when_no_pandoc(impl):
    """When pandoc is not on PATH, fallback should produce valid DOCX."""
    with patch.object(impl, "_pandoc_available", return_value=False):
        raw = impl._markdown_to_docx_bytes(_SAMPLE_MD, "fallback-test")

    doc = _load_docx(raw)
    if not doc.paragraphs:
        _fail("fallback produced empty document")
    if not doc.tables:
        _fail("fallback did not produce tables")
    _ok("fallback when no pandoc: valid docx with tables")


def _test_pandoc_path(impl):
    """If pandoc is available, verify bold + table in pandoc output."""
    if not impl._pandoc_available():
        _ok("SKIP: pandoc not available, skipping pandoc-specific tests")
        return

    raw = impl._markdown_to_docx_bytes(_SAMPLE_MD, "pandoc-test")
    doc = _load_docx(raw)

    # Check bold
    found_bold = False
    for para in doc.paragraphs:
        for run in para.runs:
            if run.bold and "加粗" in run.text:
                found_bold = True
    if not found_bold:
        _fail("pandoc path: bold not found")
    _ok("pandoc path: bold rendered")

    # Check table
    if not doc.tables:
        _fail("pandoc path: no tables found")
    t = doc.tables[0]
    _ok(f"pandoc path: table {len(t.rows)} rows x {len(t.columns)} cols")


def main() -> None:
    impl = _test_import()
    if impl is None:
        return

    _test_bold_italic(impl)
    _test_table_render(impl)
    _test_heading_font(impl)
    _test_fallback_when_no_pandoc(impl)
    _test_pandoc_path(impl)

    print("SELFTEST_PASS")


if __name__ == "__main__":
    main()
