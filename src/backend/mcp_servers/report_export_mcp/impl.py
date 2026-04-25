"""Report export implementation: DOCX and Excel."""

from __future__ import annotations

import contextlib
import io
import logging
import os
import pathlib
import re
import shutil
import subprocess
import tempfile
from typing import Any, Dict, List, Optional

from artifacts.store import save_artifact_bytes

logger = logging.getLogger(__name__)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
EXCEL_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

# 方正字体系列（body / heading），与 resources/fonts/ 中的文件对应
_BODY_FONT = "方正仿宋简体"
_HEADING_FONT = "方正小标宋简体"
_CODE_FONT = "Courier New"

_REFERENCE_DOCX = os.path.join(os.path.dirname(__file__), "reference.docx")

# Table styling constants (shared by DOCX and Excel exports)
_TABLE_HEADER_COLOR = "366092"
_TABLE_ALT_ROW_COLOR = "DCE6F1"
_TABLE_BORDER_COLOR = "B8CCE4"

# Pre-compiled regex for inline markdown parsing
_INLINE_RE = re.compile(
    r"(\*\*(.+?)\*\*)"        # group 1,2: bold
    r"|(\*(.+?)\*)"           # group 3,4: italic
    r"|(`(.+?)`)"             # group 5,6: inline code
    r"|(\[(.+?)\]\((.+?)\))"  # group 7,8,9: link [text](url)
)


# ── DOCX helpers ────────────────────────────────────────────────────────────

def _set_document_default_fonts(doc, ea_font: str) -> None:
    """Write docDefaults into styles.xml so every paragraph inherits the CJK font."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    styles_el = doc.styles.element
    docDefaults = styles_el.find(qn("w:docDefaults"))
    if docDefaults is None:
        docDefaults = OxmlElement("w:docDefaults")
        styles_el.insert(0, docDefaults)

    rPrDefault = docDefaults.find(qn("w:rPrDefault"))
    if rPrDefault is None:
        rPrDefault = OxmlElement("w:rPrDefault")
        docDefaults.append(rPrDefault)

    rPr = rPrDefault.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPrDefault.append(rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)

    rFonts.set(qn("w:eastAsia"), ea_font)
    rFonts.set(qn("w:cs"), ea_font)


def _apply_cjk_font_to_para(para, font_name: str) -> None:
    """Apply CJK font to all runs + paragraph mark in a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_rFonts(rPr_el: Any, name: str) -> None:
        rFonts = rPr_el.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr_el.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), name)
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)

    for run in para.runs:
        run.font.name = font_name
        _set_rFonts(run._element.get_or_add_rPr(), font_name)

    # paragraph mark rPr (inside pPr) so inherited font is correct
    pPr = para._element.get_or_add_pPr()
    pRpr = pPr.find(qn("w:rPr"))
    if pRpr is None:
        pRpr = OxmlElement("w:rPr")
        pPr.append(pRpr)
    _set_rFonts(pRpr, font_name)


def _setup_heading_styles(doc) -> None:
    """Set CJK font and force black color on heading style definitions."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            style = doc.styles[style_name]
            # Set font
            style.font.name = _HEADING_FONT
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), _HEADING_FONT)
            # Force black color, remove theme overrides
            style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            color_el = rPr.find(qn("w:color"))
            if color_el is not None:
                color_el.set(qn("w:val"), "000000")
                for attr in (qn("w:themeColor"), qn("w:themeShade"), qn("w:themeTint")):
                    if attr in color_el.attrib:
                        del color_el.attrib[attr]
        except KeyError:
            pass


def _pandoc_available() -> bool:
    """Check if pandoc is on PATH."""
    return shutil.which("pandoc") is not None


def _pandoc_convert(markdown: str, title: str) -> bytes:
    """Convert Markdown to DOCX bytes via pandoc subprocess."""
    # Only prepend title if the markdown doesn't already start with a heading
    first_line = markdown.lstrip().split("\n")[0] if markdown.strip() else ""
    already_has_title = first_line.startswith("#")
    full_md = markdown if (not title or already_has_title) else f"# {title}\n\n{markdown}"

    with tempfile.NamedTemporaryFile(
        suffix=".md", mode="w", encoding="utf-8", delete=False
    ) as f_in:
        f_in.write(full_md)
        in_path = f_in.name

    out_path = str(pathlib.Path(in_path).with_suffix(".docx"))
    try:
        cmd = [
            "pandoc", in_path,
            "-f", "gfm",
            "-t", "docx",
            "--wrap=none",
            "-o", out_path,
        ]
        if os.path.exists(_REFERENCE_DOCX):
            cmd += ["--reference-doc", _REFERENCE_DOCX]

        subprocess.run(cmd, check=True, capture_output=True, timeout=30)

        with open(out_path, "rb") as f:
            return f.read()
    finally:
        for p in (in_path, out_path):
            with contextlib.suppress(OSError):
                os.unlink(p)


def _post_process_cjk_fonts(docx_bytes: bytes) -> bytes:
    """Open pandoc-generated docx and ensure CJK fonts and black headings are applied."""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    _set_document_default_fonts(doc, _BODY_FONT)
    _setup_heading_styles(doc)

    for para in doc.paragraphs:
        font = _HEADING_FONT if para.style.name.startswith("Heading") else _BODY_FONT
        _apply_cjk_font_to_para(para, font)

    _style_tables(doc.tables)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _style_tables(tables) -> None:
    """Add header background, alternating row colors, and borders to tables."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor

    def _set_cell_shading(cell, color: str) -> None:
        tc_pr = cell._element.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), color)
        tc_pr.append(shading)

    def _set_cell_borders(cell) -> None:
        tc_pr = cell._element.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), _TABLE_BORDER_COLOR)
            borders.append(el)
        tc_pr.append(borders)

    for table in tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                _set_cell_borders(cell)
                if row_idx == 0:
                    _set_cell_shading(cell, _TABLE_HEADER_COLOR)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                elif row_idx % 2 == 0:
                    _set_cell_shading(cell, _TABLE_ALT_ROW_COLOR)

                for para in cell.paragraphs:
                    _apply_cjk_font_to_para(para, _BODY_FONT)


def _add_inline_runs(paragraph, text: str, font_name: str) -> None:
    """Parse inline markdown (bold, italic, inline code, links) and add runs."""
    from docx.shared import RGBColor

    last_end = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > last_end:
            run = paragraph.add_run(text[last_end:m.start()])
            run.font.name = font_name

        if m.group(2):  # bold
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.font.name = font_name
        elif m.group(4):  # italic
            run = paragraph.add_run(m.group(4))
            run.italic = True
            run.font.name = font_name
        elif m.group(6):  # inline code
            run = paragraph.add_run(m.group(6))
            run.font.name = _CODE_FONT
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif m.group(8):  # link
            run = paragraph.add_run(m.group(8))
            run.font.name = font_name
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True

        last_end = m.end()

    # Remaining plain text
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = font_name


def _table_lines_to_rows(raw_lines: List[str]) -> List[List[str]]:
    """Parse raw markdown table lines into rows of cell strings, skipping separators."""
    rows: List[List[str]] = []
    for line in raw_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        if _is_separator_row(cells):
            continue
        rows.append(cells)
    return rows


def _fallback_markdown_to_docx(markdown: str, title: str) -> bytes:
    """Pure python-docx fallback when pandoc is unavailable.

    Supports headings, lists, code blocks, basic inline formatting, and tables.
    """
    from docx import Document

    doc = Document()
    _set_document_default_fonts(doc, _BODY_FONT)
    _setup_heading_styles(doc)

    lines = (markdown or "").splitlines()

    # Only add title heading if the markdown doesn't already start with a heading
    first_non_empty = next((l.strip() for l in lines if l.strip()), "")
    if not first_non_empty.startswith("#"):
        title_para = doc.add_heading(title or "报告", level=0)
        _apply_cjk_font_to_para(title_para, _HEADING_FONT)
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        # Code block
        if stripped.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                p = doc.add_paragraph(lines[i].rstrip("\n"))
                for run in p.runs:
                    run.font.name = _CODE_FONT
                i += 1
            i += 1  # skip closing ```
            continue

        # Table detection
        if (
            stripped.startswith("|")
            and stripped.endswith("|")
            and stripped.count("|") >= 2
        ):
            table_lines = []
            while i < len(lines):
                s = lines[i].strip()
                if s.startswith("|") and s.endswith("|") and s.count("|") >= 2:
                    table_lines.append(s)
                    i += 1
                else:
                    break
            _add_table_from_lines(doc, table_lines)
            continue

        # Headings
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:].strip(), level=3)
            _apply_cjk_font_to_para(p, _HEADING_FONT)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:].strip(), level=2)
            _apply_cjk_font_to_para(p, _HEADING_FONT)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=1)
            _apply_cjk_font_to_para(p, _HEADING_FONT)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, stripped[2:].strip(), _BODY_FONT)
            _apply_cjk_font_to_para(p, _BODY_FONT)
        elif re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+\. ", "", stripped), _BODY_FONT)
            _apply_cjk_font_to_para(p, _BODY_FONT)
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, line, _BODY_FONT)
            _apply_cjk_font_to_para(p, _BODY_FONT)

        i += 1

    # Style all tables once at the end (avoids O(N²) re-styling)
    _style_tables(doc.tables)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_table_from_lines(doc, table_lines: List[str]) -> None:
    """Parse markdown table lines and add a Word table to the document."""
    rows = _table_lines_to_rows(table_lines)
    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)

    for row_idx, row_cells in enumerate(rows):
        for col_idx in range(max_cols):
            cell = table.cell(row_idx, col_idx)
            value = row_cells[col_idx] if col_idx < len(row_cells) else ""
            cell.text = ""
            para = cell.paragraphs[0]
            _add_inline_runs(para, value, _BODY_FONT)
            _apply_cjk_font_to_para(para, _BODY_FONT)


def _markdown_to_docx_bytes(markdown: str, title: str) -> bytes:
    """Convert Markdown to DOCX bytes. Uses pandoc if available, falls back to python-docx."""
    if _pandoc_available():
        logger.info("Using pandoc for markdown-to-docx conversion")
        raw = _pandoc_convert(markdown, title)
        return _post_process_cjk_fonts(raw)
    else:
        logger.info("Pandoc not available, using fallback python-docx converter")
        return _fallback_markdown_to_docx(markdown, title)


def export_report_to_docx(
    *,
    markdown: str,
    title: str = "报告",
    filename: Optional[str] = None,
    language: str = "zh",
) -> Dict[str, Any]:
    """Export markdown report to docx and return artifact metadata."""
    _ = language
    try:
        blob = _markdown_to_docx_bytes(markdown=markdown, title=title)
    except Exception as e:
        return {"ok": False, "error": str(e)}

    base_name = (filename or title or "报告").strip() or "报告"
    if not base_name.lower().endswith(".docx"):
        base_name = f"{base_name}.docx"

    item = save_artifact_bytes(
        content=blob,
        name=base_name,
        mime_type=DOCX_MIME,
        extension="docx",
        metadata={"title": title},
    )

    return {
        "ok": True,
        "file_id": item["file_id"],
        "name": item["name"],
        "mime_type": item["mime_type"],
        "size": item["size"],
        "url": f"/files/{item['file_id']}",
    }


# ── Excel helpers ────────────────────────────────────────────────────────────

def _is_separator_row(cells: List[str]) -> bool:
    """Return True if every non-empty cell looks like |---|:---:|---| separator."""
    non_empty = [c.strip() for c in cells if c.strip()]
    if not non_empty:
        return False
    return all(bool(re.match(r"^:?-+:?$", c)) for c in non_empty)


def _parse_markdown_tables(text: str) -> List[List[List[str]]]:
    """
    Extract all markdown tables from *text*.
    Returns a list of tables; each table is a list of rows (list of cell strings).
    Separator rows (|---|---|) are excluded; the first row becomes the header.
    """
    lines = text.splitlines()
    raw_tables: List[List[str]] = []
    current: List[str] = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2:
            current.append(stripped)
        else:
            if current:
                raw_tables.append(current)
                current = []
    if current:
        raw_tables.append(current)

    result: List[List[List[str]]] = []
    for raw_table in raw_tables:
        rows = _table_lines_to_rows(raw_table)
        if rows:
            result.append(rows)
    return result


def export_table_to_excel(
    *,
    markdown: str,
    title: str = "表格",
    filename: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Parse markdown table(s) from *markdown* and export them as an Excel workbook.
    Each markdown table becomes one sheet.  Returns artifact metadata on success.
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError as e:
        return {"ok": False, "error": "openpyxl is not installed; add it to requirements.txt"}

    tables = _parse_markdown_tables(markdown)
    if not tables:
        return {"ok": False, "error": "未在输入中找到 Markdown 表格，请确保内容包含标准 | 列 | 格式 | 表格"}

    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # remove default empty sheet

    # Shared styles
    header_fill = PatternFill(start_color=_TABLE_HEADER_COLOR, end_color=_TABLE_HEADER_COLOR, fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    data_font = Font(size=11)
    alt_fill = PatternFill(start_color=_TABLE_ALT_ROW_COLOR, end_color=_TABLE_ALT_ROW_COLOR, fill_type="solid")
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
    thin = Side(style="thin", color=_TABLE_BORDER_COLOR)
    cell_border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for idx, rows in enumerate(tables):
        # Sheet name: use title for single table, Sheet1/2/... for multiples
        if len(tables) == 1:
            sheet_name = (title or "数据")[:31]
        else:
            sheet_name = f"Sheet{idx + 1}"
        ws = wb.create_sheet(title=sheet_name)

        max_cols = max(len(r) for r in rows)

        for row_idx, cells in enumerate(rows, start=1):
            is_header = row_idx == 1
            is_alt_row = (not is_header) and (row_idx % 2 == 0)

            for col_idx in range(1, max_cols + 1):
                value = cells[col_idx - 1] if col_idx <= len(cells) else ""
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border = cell_border
                if is_header:
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = center_align
                else:
                    cell.font = data_font
                    cell.alignment = left_align
                    if is_alt_row:
                        cell.fill = alt_fill

        # Auto-fit column widths (estimate by character length)
        for col in ws.columns:
            max_len = 0
            for cell in col:
                try:
                    text = str(cell.value or "")
                    # CJK chars ≈ 2 half-width chars wide
                    w = sum(2 if ord(c) > 0x2E80 else 1 for c in text)
                    max_len = max(max_len, w)
                except Exception:
                    pass
            ws.column_dimensions[col[0].column_letter].width = min(max(max_len + 2, 10), 60)

        # Freeze header row
        ws.freeze_panes = "A2"

    buf = io.BytesIO()
    wb.save(buf)
    blob = buf.getvalue()

    base_name = (filename or title or "表格").strip() or "表格"
    if not base_name.lower().endswith(".xlsx"):
        base_name = f"{base_name}.xlsx"

    item = save_artifact_bytes(
        content=blob,
        name=base_name,
        mime_type=EXCEL_MIME,
        extension="xlsx",
        metadata={"title": title, "sheet_count": len(tables)},
    )

    return {
        "ok": True,
        "file_id": item["file_id"],
        "name": item["name"],
        "mime_type": item["mime_type"],
        "size": item["size"],
        "url": f"/files/{item['file_id']}",
        "sheet_count": len(tables),
    }
